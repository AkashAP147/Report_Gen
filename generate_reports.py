import pandas as pd
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def extract_subjects(excel_file):
    try:
        tables = pd.read_html(excel_file)
        if not tables: return []
        df = tables[0]
        if 'S No' not in df.columns and 'USN' not in df.columns:
            df.columns = df.iloc[0]
            df = df[1:].reset_index(drop=True)
        df = df.dropna(how='all')
        if 'Subject Code' in df.columns:
            return df['Subject Code'].dropna().unique().tolist()
        return []
    except Exception as e:
        print(f"Extraction error: {e}")
        return []

def assign_blocks(df, max_per_block=32, pref_left=None, pref_right=None):
    if 'Branch' not in df.columns:
        df['Branch'] = 'ALL'
    subject_students = {subj: group.to_dict('records') for subj, group in df.groupby(['Subject Code', 'Branch'])}
    arranged_students = []
    
    block_labels = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII', 'XIII', 'XIV', 'XV', 'XVI', 'XVII', 'XVIII', 'XIX', 'XX']
    block_idx = 0
    
    while subject_students:
        block_no_str = f"Block-{block_labels[block_idx] if block_idx < len(block_labels) else block_idx + 1}"
        
        # Pick largest subject to start the left side, prioritizing preferred subjects
        def get_sort_key(k):
            boost = 0
            if pref_left and k[0] == pref_left: boost = 1000000
            elif pref_right and k[0] == pref_right: boost = 500000
            return len(subject_students[k]) + boost
            
        sorted_subjs = sorted(subject_students.keys(), key=get_sort_key, reverse=True)
        subj_left = sorted_subjs[0]
        
        # Pick second largest subject to start the right side
        subj_right = sorted_subjs[1] if len(sorted_subjs) > 1 else None
        
        single_subj_turn = 'left'
        
        # Fill row by row to ensure even depth on both sides
        for i in range(16):
            if not subject_students:
                break
                
            desk_odd = i * 2 + 1
            desk_even = i * 2 + 2
            
            # If only one subject remains globally, zig-zag it to balance the columns evenly
            if len(subject_students) == 1:
                subj = list(subject_students.keys())[0]
                if single_subj_turn == 'left':
                    student = subject_students[subj].pop(0)
                    student['Block No'] = block_no_str
                    student['Desk No'] = desk_odd
                    arranged_students.append(student)
                    single_subj_turn = 'right'
                else:
                    student = subject_students[subj].pop(0)
                    student['Block No'] = block_no_str
                    student['Desk No'] = desk_even
                    arranged_students.append(student)
                    single_subj_turn = 'left'
                    
                if not subject_students[subj]:
                    del subject_students[subj]
                continue
            
            # Fill Odd Desk (Left Column)
            if subj_left not in subject_students:
                subj_left = None
                    
            if subj_left is not None:
                student = subject_students[subj_left].pop(0)
                student['Block No'] = block_no_str
                student['Desk No'] = desk_odd
                arranged_students.append(student)
                
                if not subject_students[subj_left]:
                    del subject_students[subj_left]
                    
            # Fill Even Desk (Right Column)
            if subj_right not in subject_students:
                subj_right = None
                    
            if subj_right is not None:
                student = subject_students[subj_right].pop(0)
                student['Block No'] = block_no_str
                student['Desk No'] = desk_even
                arranged_students.append(student)
                
                if not subject_students[subj_right]:
                    del subject_students[subj_right]
                    
        block_idx += 1
        
    return pd.DataFrame(arranged_students)

def parse_excel(filepath, pref_left=None, pref_right=None):
    # Read the html file
    dfs = pd.read_html(filepath)
    
    # The first table is likely the header
    # But usually pd.read_html just gets the big table.
    # Let's read the raw html to get header info
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        html_content = f.read()
        
    import re
    college_code = 'VS'
    date = '2026-09-03'
    session = '09:30:00'
    
    if 'CollegeCode -' in html_content:
        m = re.search(r'CollegeCode - ([\w]+)', html_content)
        if m: college_code = m.group(1).strip()
    if 'Date -' in html_content:
        m = re.search(r'Date - ([\d\-]+)', html_content)
        if m: date = m.group(1).strip()
    if 'Session -' in html_content:
        m = re.search(r'Session - ([\d:]+)', html_content)
        if m: session = m.group(1).strip()

    # The actual data table is usually the largest one or the only one parsed
    df_data = dfs[0]
    # Clean up columns if it's the main table
    if 'S No' in df_data.columns or 'USN' in df_data.columns:
        pass # Already has headers
    else:
        # If headers are in row 0
        df_data.columns = df_data.iloc[0]
        df_data = df_data[1:].reset_index(drop=True)
        
    def extract_branch(usn):
        m = re.search(r'\d[A-Z]{2}\d{2}([A-Z]{2,3})', str(usn).upper())
        if m:
            return m.group(1)
        return "ALL"
        
    if 'USN' in df_data.columns:
        df_data['Branch'] = df_data['USN'].apply(extract_branch)
    else:
        df_data['Branch'] = "ALL"
        
    df_data = assign_blocks(df_data, pref_left=pref_left, pref_right=pref_right)
        
    return {
        'college_code': college_code,
        'date': date,
        'session': session,
        'data': df_data
    }

def create_block_list(data_dict, output_dir):
    df = data_dict['data']
    date = data_dict['date']
    session = data_dict['session']
    
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re
    import os
    
    # Format Time
    if "09:" in session or "10:" in session or "11:" in session or "AM" in session.upper():
        formatted_time = "09.30am to 12.30pm"
    else:
        formatted_time = "02.00pm to 05.00pm"
        
    date_parts = date.split('-')
    if len(date_parts) == 3:
        formatted_date = f"{date_parts[2]}.{date_parts[1]}.{date_parts[0]}"
    else:
        formatted_date = date
    
    doc = docx.Document()
    
    # Set narrow margins
    section = doc.sections[-1]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    for block_no, group in df.groupby('Block No', sort=False):
        def make_image_float_align(run, align='left'):
            """Convert an inline image in a run to a floating (anchored) image with alignment."""
            drawing = run._r.find(qn('w:drawing'))
            inline = drawing.find(qn('wp:inline'))
            
            anchor = OxmlElement('wp:anchor')
            anchor.set('distT', '0')
            anchor.set('distB', '0')
            anchor.set('distL', '114300')
            anchor.set('distR', '114300')
            anchor.set('simplePos', '0')
            anchor.set('relativeHeight', '251658240')
            anchor.set('behindDoc', '0')
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '1')
            anchor.set('allowOverlap', '1')
            
            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)
            
            posH = OxmlElement('wp:positionH')
            posH.set('relativeFrom', 'margin')
            align_node = OxmlElement('wp:align')
            align_node.text = align # 'left' or 'right'
            posH.append(align_node)
            anchor.append(posH)
            
            posV = OxmlElement('wp:positionV')
            posV.set('relativeFrom', 'paragraph')
            posOffset_v = OxmlElement('wp:posOffset')
            posOffset_v.text = '0'
            posV.append(posOffset_v)
            anchor.append(posV)
            
            for child in list(inline):
                anchor.append(child)
            
            wrapNone = OxmlElement('wp:wrapNone')
            anchor.insert(4, wrapNone)
            
            drawing.remove(inline)
            drawing.append(anchor)

        # Header Title Paragraph
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # VSM Logo Left
        vsm_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'vsm_logo.jpg')
        if os.path.exists(vsm_path):
            r_vsm = p_title.add_run()
            r_vsm.add_picture(vsm_path, width=Inches(0.9))
            make_image_float_align(r_vsm, 'left')
            
        # VTU Logo Right
        vtu_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'vtu_logo.png')
        if os.path.exists(vtu_path):
            r_vtu = p_title.add_run()
            r_vtu.add_picture(vtu_path, width=Inches(0.9))
            make_image_float_align(r_vtu, 'right')
            
        # Title Center Text
        r_title = p_title.add_run("VSM'S S R K INSTITUTE OF TECHNOLOGY\n")
        r_title.bold = True
        r_title.font.size = Pt(16)
        r_sub = p_title.add_run("Nipani - 591 237, Dist: Belgaum, Karnataka State")
        r_sub.bold = True
        r_sub.font.size = Pt(10)
        
        # Horizontal Line
        p_line = doc.add_paragraph()
        pPr = p_line._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        # Examination Name
        p_exam = doc.add_paragraph()
        p_exam.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_exam = p_exam.add_run("THEORY EXAMINATIONS JUNE/JULY-2026")
        r_exam.bold = True
        r_exam.underline = True
        r_exam.font.size = Pt(12)
        
        # Details section using tab stops
        from docx.enum.text import WD_TAB_ALIGNMENT
        p_details = doc.add_paragraph()
        tab_stops = p_details.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(1.2), WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.LEFT)
        
        subjects = group['Subject Code'].unique()
        subj_str = ", ".join(map(str, subjects))
        
        # Try to infer semester
        sem = ""
        for subj in subjects:
            m = re.search(r'[A-Za-z]+(\d)', str(subj))
            if m:
                sem = m.group(1)
                break
        
        subject_names = group['Subject Name'].unique()
        subj_name_str = ", ".join(map(str, subject_names))
        
        r1 = p_details.add_run("Subject:")
        r1.bold = True
        p_details.add_run(f"\t{subj_name_str}\n")
        
        r2 = p_details.add_run("Subject Code:")
        r2.bold = True
        p_details.add_run(f"\t{subj_str}\t")
        r3 = p_details.add_run("Semester:")
        r3.bold = True
        p_details.add_run(f"\t{sem}\n")
        
        r4 = p_details.add_run("Date:")
        r4.bold = True
        p_details.add_run(f"\t{formatted_date}\t")
        r5 = p_details.add_run("Time:")
        r5.bold = True
        p_details.add_run(f"\t{formatted_time}")
        
        # Block Number
        p_block = doc.add_paragraph()
        p_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_block.add_run(f"{block_no.upper().replace('-', ' - ')}")
        run.bold = True
        run.underline = True
        run.font.size = Pt(14)
        
        # Data Table
        table = doc.add_table(rows=17, cols=4)
        table.style = 'Table Grid'
        
        from docx.enum.table import WD_TABLE_ALIGNMENT
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Fix table widths
        tbl_el2 = table._tbl
        tblPr2 = tbl_el2.tblPr if tbl_el2.tblPr is not None else OxmlElement('w:tblPr')
        tblLayout2 = OxmlElement('w:tblLayout')
        tblLayout2.set(qn('w:type'), 'fixed')
        tblPr2.append(tblLayout2)
        
        for r in table.rows:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            r.cells[0].width = Inches(2.2)
            r.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r.cells[1].width = Inches(1.5)
            r.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r.cells[2].width = Inches(2.2)
            r.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r.cells[3].width = Inches(1.5)
            r.cells[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(['USN', 'DESK NO.', 'USN', 'DESK NO.']):
            hdr_cells[idx].text = text
            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[idx].paragraphs[0].runs[0].bold = True
            
        # Pre-fill desk numbers
        for i in range(1, 17):
            row = table.rows[i]
            # Left desk (Odd)
            desk_left = i * 2 - 1
            row.cells[1].text = str(desk_left)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].paragraphs[0].runs[0].bold = True
            
            # Right desk (Even)
            desk_right = i * 2
            row.cells[3].text = str(desk_right)
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[3].paragraphs[0].runs[0].bold = True
            
        # Map USNs by desk
        desk_to_usn = {row['Desk No']: row['USN'] for _, row in group.iterrows()}
        desk_to_subj = {row['Desk No']: row['Subject Code'] for _, row in group.iterrows()}
        seen_subjects = set()
        
        for i in range(1, 17):
            row = table.rows[i]
            desk_left = i * 2 - 1
            desk_right = i * 2
            
            if desk_left in desk_to_usn:
                usn_text = str(desk_to_usn[desk_left])
                subj_code = str(desk_to_subj[desk_left])
                if subj_code not in seen_subjects:
                    usn_text += f"\n({subj_code})"
                    seen_subjects.add(subj_code)
                row.cells[0].text = usn_text
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            if desk_right in desk_to_usn:
                usn_text = str(desk_to_usn[desk_right])
                subj_code = str(desk_to_subj[desk_right])
                if subj_code not in seen_subjects:
                    usn_text += f"\n({subj_code})"
                    seen_subjects.add(subj_code)
                row.cells[2].text = usn_text
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set font size 16 and line spacing 1.5 for all cells in the table
        for r in table.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    p.paragraph_format.line_spacing = 1.5
                    for run in p.runs:
                        run.font.size = Pt(16)
            
        doc.add_page_break()
        
    output_path = os.path.join(output_dir, 'Block_List_Generated.docx')
    doc.save(output_path)
    print(f"Generated: {output_path}")

def create_dispatch_format_i(data_dict, output_dir):
    df = data_dict['data']
    date = data_dict['date']
    session = data_dict['session']
    
    # Format Date
    date_parts = date.split('-')
    if len(date_parts) == 3:
        formatted_date = f"{date_parts[2]}.{date_parts[1]}.{date_parts[0]}"
    else:
        formatted_date = date
        
    # Format Time
    if "09:" in session or "10:" in session or "11:" in session or "AM" in session.upper():
        formatted_time = "09.30 AM TO 12.30 PM"
    else:
        formatted_time = "02.00 PM TO 05.00 PM"
    
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    def set_page_borders(section):
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '24')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)
        sectPr.append(pgBorders)

    doc = docx.Document()
    
    # Set Landscape for the whole document ONCE
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Standard letter portrait is 8.5 x 11. Swap to 11 x 8.5
    section.page_width = Inches(11.69) # A4 width in landscape
    section.page_height = Inches(8.27) # A4 height in landscape
    set_page_borders(section)
    
    for (subject_code, branch), group in df.groupby(['Subject Code', 'Branch']):
        
        # Helper to remove all borders from a table (make it invisible)
        def remove_table_borders(tbl):
            tbl_el = tbl._tbl
            tblPr = tbl_el.tblPr if tbl_el.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            tblPr.append(tblBorders)
        
        def disable_autofit(tbl):
            """Disable auto-fit so explicit column widths are respected."""
            tbl_el = tbl._tbl
            tblPr = tbl_el.tblPr if tbl_el.tblPr is not None else OxmlElement('w:tblPr')
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
        
        def make_image_float(run, left_emu, top_emu):
            """Convert an inline image in a run to a floating (anchored) image."""
            drawing = run._r.find(qn('w:drawing'))
            inline = drawing.find(qn('wp:inline'))
            
            # Create anchor element
            anchor = OxmlElement('wp:anchor')
            anchor.set('distT', '0')
            anchor.set('distB', '0')
            anchor.set('distL', '114300')
            anchor.set('distR', '114300')
            anchor.set('simplePos', '0')
            anchor.set('relativeHeight', '251658240')
            anchor.set('behindDoc', '0')
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '1')
            anchor.set('allowOverlap', '1')
            
            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)
            
            # Horizontal position relative to column
            posH = OxmlElement('wp:positionH')
            posH.set('relativeFrom', 'column')
            posOffset_h = OxmlElement('wp:posOffset')
            posOffset_h.text = str(left_emu)
            posH.append(posOffset_h)
            anchor.append(posH)
            
            # Vertical position relative to paragraph
            posV = OxmlElement('wp:positionV')
            posV.set('relativeFrom', 'paragraph')
            posOffset_v = OxmlElement('wp:posOffset')
            posOffset_v.text = str(top_emu)
            posV.append(posOffset_v)
            anchor.append(posV)
            
            # Move all children from inline to anchor
            for child in list(inline):
                anchor.append(child)
            
            # Add wrap none so text flows freely
            wrapNone = OxmlElement('wp:wrapNone')
            anchor.insert(4, wrapNone)
            
            # Replace inline with anchor
            drawing.remove(inline)
            drawing.append(anchor)
        
        # Add floating logo on the first header paragraph
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Insert logo as floating image
        logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'vtu_logo.png')
        if os.path.exists(logo_path):
            run_logo = p_title.add_run()
            run_logo.add_picture(logo_path, width=Inches(0.85))
            # Position: 0 EMU from left column, 0 EMU from top of paragraph
            make_image_float(run_logo, 0, 0)
        
        # Header text as free centered runs on the same paragraph
        run_title = p_title.add_run("Book Dispatch Format – I\n")
        run_title.bold = True
        run_title.font.size = Pt(14)
        
        run_univ = p_title.add_run("VISVESVARAYA TECHNOLOGICAL UNIVERSITY, BELGAUM\n")
        run_univ.bold = True
        run_univ.font.size = Pt(14)
        
        run_region = p_title.add_run("BELGAUM REGION\n")
        run_region.bold = True
        run_region.font.size = Pt(14)
        
        run_sem = p_title.add_run("B.E- V Semester Examination JUNE/JULY-2026")
        run_sem.bold = True
        run_sem.font.size = Pt(14)
        
        # Extract branch from USN
        usn = str(group['USN'].iloc[0]).upper()
        m = re.search(r'\d[A-Z]{2}\d{2}([A-Z]{2,3})', usn)
        if m:
            enclosure = m.group(1)
        else:
            m2 = re.search(r'[A-Za-z]+', str(subject_code))
            enclosure = m2.group(0) if m2 else "ALL"
        
        # Exam Centre line
        p_centre = doc.add_paragraph()
        run_ec = p_centre.add_run("Exam Centre: - V.S.M's S.R.K Institute of Technology, Nipani")
        run_ec.bold = True
        run_ec.font.size = Pt(11)
        
        # Enclosures line
        p_enc = doc.add_paragraph()
        run_enc = p_enc.add_run(f"Enclosures: - {enclosure}")
        run_enc.bold = True
        run_enc.font.size = Pt(11)
        
        p_asd = doc.add_paragraph()
        run_asd = p_asd.add_run("Answer Script Details:-")
        run_asd.bold = True
        run_asd.font.size = Pt(11)
        
        # Data table - centered, compact, fixed layout
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Cm
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        disable_autofit(table)
        
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(2.8)
        table.columns[2].width = Inches(2.0)
        table.columns[3].width = Inches(1.7)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date of Exam'
        hdr_cells[1].text = 'Exam Timing'
        hdr_cells[2].text = 'Subject Code'
        hdr_cells[3].text = 'No. of Papers'
        
        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(14)
        
        row_cells = table.add_row().cells
        row_cells[0].text = formatted_date
        row_cells[1].text = formatted_time
        row_cells[2].text = str(subject_code)
        row_cells[3].text = f"{enclosure}-{len(group):02d}"
        
        # Set data row height and bold large font, vertically center
        table.rows[1].height = Cm(1.5)
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Forms check line
        p_forms = doc.add_paragraph()
        p_forms.add_run("\n")
        run_f = p_forms.add_run("2. Form –A   \u2714    \t\t3. Form- B   \u2714  \t\t4. Question Paper   \u2714")
        run_f.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # To section
        p_to = doc.add_paragraph()
        p_to.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_to1 = p_to.add_run("To,\n")
        run_to1.bold = True
        run_to1.font.size = Pt(12)
        run_to2 = p_to.add_run("Dr. Sattagouda M Patil\n")
        run_to2.bold = True
        run_to2.font.size = Pt(12)
        run_to3 = p_to.add_run("Chief Co-ordinator\n")
        run_to3.bold = True
        run_to3.font.size = Pt(12)
        run_to4 = p_to.add_run('VTU, Digitization Centre  "Jnana Sangama"\n')
        run_to4.bold = True
        run_to4.font.size = Pt(12)
        run_to5 = p_to.add_run("VTU, Belagavi : 590018")
        run_to5.bold = True
        run_to5.font.size = Pt(12)
        
        # From section
        p_from = doc.add_paragraph()
        run_from1 = p_from.add_run("From,\n")
        run_from1.bold = True
        run_from1.font.size = Pt(12)
        run_from2 = p_from.add_run("       Chief Superintendent")
        run_from2.font.size = Pt(11)
        
        doc.add_page_break()
        
    output_path = os.path.join(output_dir, 'Dispatch_Format_I_Generated.docx')
    doc.save(output_path)
    print(f"Generated: {output_path}")

def create_dispatch_format_ii(data_dict, output_dir):
    df = data_dict['data']
    date = data_dict['date']
    session = data_dict['session']
    
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    # Format date
    date_parts = date.split('-')
    if len(date_parts) == 3:
        formatted_date = f"{date_parts[2]}.{date_parts[1]}.{date_parts[0]}"
    else:
        formatted_date = date
    
    session_str = "MOR" if ("09:" in session or "10:" in session or "11:" in session or "AM" in session.upper()) else "AFT"
    
    def make_image_float(run, left_emu, top_emu):
        """Convert an inline image in a run to a floating (anchored) image."""
        drawing = run._r.find(qn('w:drawing'))
        inline = drawing.find(qn('wp:inline'))
        
        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '114300')
        anchor.set('distR', '114300')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251658240')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')
        
        simplePos = OxmlElement('wp:simplePos')
        simplePos.set('x', '0')
        simplePos.set('y', '0')
        anchor.append(simplePos)
        
        posH = OxmlElement('wp:positionH')
        posH.set('relativeFrom', 'column')
        posOffset_h = OxmlElement('wp:posOffset')
        posOffset_h.text = str(left_emu)
        posH.append(posOffset_h)
        anchor.append(posH)
        
        posV = OxmlElement('wp:positionV')
        posV.set('relativeFrom', 'paragraph')
        posOffset_v = OxmlElement('wp:posOffset')
        posOffset_v.text = str(top_emu)
        posV.append(posOffset_v)
        anchor.append(posV)
        
        for child in list(inline):
            anchor.append(child)
        
        wrapNone = OxmlElement('wp:wrapNone')
        anchor.insert(4, wrapNone)
        
        drawing.remove(inline)
        drawing.append(anchor)
    
    def disable_autofit(tbl):
        tbl_el = tbl._tbl
        tblPr = tbl_el.tblPr if tbl_el.tblPr is not None else OxmlElement('w:tblPr')
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
    
    doc = docx.Document()
    
    # Set narrow margins
    section = doc.sections[-1]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Header with floating logo
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'vtu_logo.png')
    if os.path.exists(logo_path):
        run_logo = p_title.add_run()
        run_logo.add_picture(logo_path, width=Inches(0.7))
        make_image_float(run_logo, 0, 0)
    
    r1 = p_title.add_run("Answer Book Dispatch Format – II\n")
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(11)
    
    r2 = p_title.add_run("Visvesvaraya Technological University, Belagavi\n")
    r2.bold = True
    r2.font.size = Pt(13)
    
    r3 = p_title.add_run("BELAGAVI REGION\n")
    r3.bold = True
    r3.font.size = Pt(9)
    
    r4 = p_title.add_run("ANSWER BOOK BUNDLES ACKNOWLEDGEMENT\n")
    r4.bold = True
    r4.font.size = Pt(9)
    
    r5 = p_title.add_run("JUNE/JULY-2026 Examinations")
    r5.bold = True
    r5.font.size = Pt(9)
    
    # Date line - right aligned
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = p_date.add_run(f"Date: {formatted_date}")
    run_date.font.size = Pt(9)
    
    # Exam Centre
    p_ec = doc.add_paragraph()
    run_ec = p_ec.add_run("Exam Centre: VSM INSTITUTE OF TECHNOLOGY, NIPANI.")
    run_ec.font.size = Pt(9)
    
    # Valuation Centre
    p_vc = doc.add_paragraph()
    run_vc = p_vc.add_run('Valuation Centre: Prof. Sattagouda M Patil ,Chief Co-ordinator, VTU, Digitization Centre "Jnana Sangama" VTU, Belagavi-18')
    run_vc.font.size = Pt(9)
    
    # Data table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    disable_autofit(table)
    
    table.columns[0].width = Cm(1.0)
    table.columns[1].width = Cm(2.2)
    table.columns[2].width = Cm(2.2)
    table.columns[3].width = Cm(8.0)
    table.columns[4].width = Cm(2.2)
    table.columns[5].width = Cm(2.2)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sl.\nNo.'
    hdr_cells[1].text = 'Date of\nExam'
    hdr_cells[2].text = 'Subject\nCode'
    hdr_cells[3].text = 'Subject Title'
    hdr_cells[4].text = 'No. of\nAns.\nScripts'
    hdr_cells[5].text = 'No. of\nBundles'
    
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    sl_no = 1
    total_bundles = 0
    
    for (subject_code, branch), group in df.groupby(['Subject Code', 'Branch']):
        subject_name = group['Subject Name'].iloc[0]
        scripts_count = len(group)
        bundles_count = (scripts_count // 20) + 1
        
        # Extract branch from USN
        usn = str(group['USN'].iloc[0]).upper()
        m = re.search(r'\d[A-Z]{2}\d{2}([A-Z]{2,3})', usn)
        if m:
            enclosure = m.group(1)
        else:
            m2 = re.search(r'[A-Za-z]+', str(subject_code))
            enclosure = m2.group(0) if m2 else "ALL"
        
        row_cells = table.add_row().cells
        row_cells[0].text = f"{sl_no:02d}"
        
        # We will merge the Date of Exam cell later, but populate it first
        row_cells[1].text = f"{formatted_date}\n{session_str}"
        row_cells[2].text = str(subject_code)
        row_cells[3].text = str(subject_name)
        row_cells[4].text = f"{enclosure}-{scripts_count:02d}"
        row_cells[5].text = f"{bundles_count:02d}"
        
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        total_bundles += bundles_count
        sl_no += 1
    
    # Merge 'Date of Exam' cells (Column 1) for all data rows
    if len(table.rows) > 1:
        start_cell = table.cell(1, 1)
        end_cell = table.cell(len(table.rows) - 1, 1)
        if start_cell != end_cell:
            start_cell.merge(end_cell)
            start_cell.text = f"{formatted_date}\n{session_str}"
            start_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in start_cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    
    # Add TOTAL row
    row_cells = table.add_row().cells
    cell_start = row_cells[0]
    cell_end = row_cells[4]
    cell_start.merge(cell_end)
    cell_start.text = "TOTAL"
    p = cell_start.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(10)
    
    row_cells[5].text = f"{total_bundles:02d}"
    p = row_cells[5].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(10)
    
    # Force widths on all cells to prevent cutting
    widths = [Inches(0.4), Inches(0.8), Inches(0.8), Inches(3.5), Inches(0.8), Inches(0.8)]
    for row in table.rows:
        # Check if row has 6 cells (not the merged TOTAL row)
        if len(row.cells) == 6:
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]
    
    # Set alignment for all cells
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Helper to convert numbers to words
    def num2words(n):
        words = {1: 'ONE', 2: 'TWO', 3: 'THREE', 4: 'FOUR', 5: 'FIVE', 6: 'SIX', 7: 'SEVEN', 8: 'EIGHT', 9: 'NINE', 10: 'TEN',
                 11: 'ELEVEN', 12: 'TWELVE', 13: 'THIRTEEN', 14: 'FOURTEEN', 15: 'FIFTEEN', 16: 'SIXTEEN', 17: 'SEVENTEEN',
                 18: 'EIGHTEEN', 19: 'NINETEEN', 20: 'TWENTY', 30: 'THIRTY', 40: 'FORTY', 50: 'FIFTY', 60: 'SIXTY', 70: 'SEVENTY', 80: 'EIGHTY', 90: 'NINETY'}
        if n in words: return words[n]
        if n < 100:
            tens, units = divmod(n, 10)
            return words[tens*10] + (" " + words[units] if units else "")
        return str(n)
        
    p_recv = doc.add_paragraph()
    p_recv.add_run("\nReceived (")
    r_num = p_recv.add_run(num2words(total_bundles))
    r_num.bold = True
    r_num.underline = True
    p_recv.add_run(") Sealed Answer book bundles.\n\n\n")
    
    # Signatures - use free text with tab stops for precise positioning
    from docx.enum.text import WD_TAB_ALIGNMENT
    p_sig = doc.add_paragraph()
    tab_stops = p_sig.paragraph_format.tab_stops
    # Add center-aligned tab stops for the left and right signature blocks
    tab_stops.add_tab_stop(Inches(1.5), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.CENTER)
    
    p_sig.add_run("\tChief Superintendent\tMember, Collection Team\n")
    p_sig.add_run("\t(Signature with date & Seal)\t(Signature with date)")
    
    output_path = os.path.join(output_dir, 'Dispatch_Format_II_Generated.docx')
    doc.save(output_path)
    print(f"Generated: {output_path}")

def create_consolidated_list(data_dict, output_dir):
    df = data_dict['data']
    date = data_dict['date']
    session = data_dict['session']
    
    from docx.shared import Cm, Inches, Pt
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    
    doc = docx.Document()
    
    # Set Narrow Margins to fit the table better
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VSM INSTITUTE OF TECHNOLOGY, NIPANI\nVTU THEORY EXAMINATION JUNE/JULY-2026\nConsolidated block wise information")
    run.bold = True
    
    # Add Time and Date
    total_candidates = len(df)
    p2 = doc.add_paragraph(f"TIME :- {session} \t\t\t\t DATE: {date} \t\t\t\t TOTAL: {total_candidates:02d}")
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.allow_autofit = False
    
    hdr_cells = table.rows[0].cells
    headers = ['Block No', 'Subject Code', 'Subject Name', 'USN Numbers', 'Total No. Of Candidates', 'Invigilator\nSign']
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        for p_hdr in hdr_cells[idx].paragraphs:
            p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_hdr in p_hdr.runs:
                run_hdr.bold = True
                run_hdr.font.size = Pt(10)
    
    # Adjusted widths to fit inside narrow margins (approx 19 cm total width available)
    widths = [Cm(1.5), Cm(2.0), Cm(2.8), Cm(9.5), Cm(1.7), Cm(1.5)]
    
    for block_no, block_group in df.groupby('Block No', sort=False):
        start_block_row_idx = len(table.rows)
        
        for subject_code, subject_group in block_group.groupby('Subject Code'):
            start_subj_row_idx = len(table.rows)
            
            for branch, group in subject_group.groupby('Branch'):
                subject_name = group['Subject Name'].iloc[0]
                usns = group['USN'].tolist()
                
                # Group into chunks of 4 and join with newlines (NO SPACES AFTER COMMAS)
                usns_chunks = [usns[i:i + 4] for i in range(0, len(usns), 4)]
                usns_str = "\n".join([",".join(map(str, chunk)) for chunk in usns_chunks])
                
                row_cells = table.add_row().cells
                row_cells[0].text = block_no
                row_cells[0].paragraphs[0].runs[0].bold = True
                
                row_cells[1].text = str(subject_code)
                row_cells[1].paragraphs[0].runs[0].bold = True
                
                row_cells[2].text = str(subject_name)
                row_cells[2].paragraphs[0].runs[0].bold = True
                
                row_cells[3].text = usns_str
                row_cells[4].text = f"{len(usns)}"
                row_cells[5].text = ''
                
                # Set font size for all cells in this row to 10pt
                for cell in row_cells:
                    for p_cell in cell.paragraphs:
                        for run_cell in p_cell.runs:
                            run_cell.font.size = Pt(10)
                            
            # Merge Subject Code and Subject Name cells if multiple branches
            end_subj_row_idx = len(table.rows) - 1
            if end_subj_row_idx > start_subj_row_idx:
                table.cell(start_subj_row_idx, 1).merge(table.cell(end_subj_row_idx, 1))
                table.cell(start_subj_row_idx, 1).text = str(subject_code)
                table.cell(start_subj_row_idx, 1).paragraphs[0].runs[0].bold = True
                
                table.cell(start_subj_row_idx, 2).merge(table.cell(end_subj_row_idx, 2))
                table.cell(start_subj_row_idx, 2).text = str(subject_name)
                table.cell(start_subj_row_idx, 2).paragraphs[0].runs[0].bold = True
            
        # Merge block_no and invigilator sign cells if multiple subjects/branches
        end_block_row_idx = len(table.rows) - 1
        if end_block_row_idx > start_block_row_idx:
            # Merge Block No (col 0)
            table.cell(start_block_row_idx, 0).merge(table.cell(end_block_row_idx, 0))
            table.cell(start_block_row_idx, 0).text = block_no
            table.cell(start_block_row_idx, 0).paragraphs[0].runs[0].bold = True
            
            # Merge Invigilator Sign (col 5)
            table.cell(start_block_row_idx, 5).merge(table.cell(end_block_row_idx, 5))
            
    # Apply widths, alignments, and vertical centering
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph("\n\n\nDeputy Chief Supdt. \t\t\t\t\t\t\t Chief Supdt.")
    
    output_path = os.path.join(output_dir, 'Consolidated_List_Generated.docx')
    doc.save(output_path)
    print(f"Generated: {output_path}")

def run_generation(excel_file, output_dir, pref_left=None, pref_right=None):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    print(f"Reading {excel_file}...")
    try:
        data = parse_excel(excel_file, pref_left, pref_right)
        print("Data loaded successfully. Generating reports...")
        
        create_block_list(data, output_dir)
        create_dispatch_format_i(data, output_dir)
        create_dispatch_format_ii(data, output_dir)
        create_consolidated_list(data, output_dir)
        
        print(f"All reports generated successfully in {output_dir}")
        return True
    except Exception as e:
        print(f"Error occurred: {e}")
        return False

if __name__ == '__main__':
    base_dir = r"d:\COMMERCIAL\DOC_COV"
    excel_file = os.path.join(base_dir, "VS_2026-09-03_09_30_00_Session.xls")
    output_dir = os.path.join(base_dir, "Generated_Reports")
    run_generation(excel_file, output_dir)
